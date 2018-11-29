#coding=utf-8
#!/usr/bin/python3
import docx
import sys
import string
import os
import ListFolder
from docx.shared import Inches

#读取指定的文档，
# 2 读取指定目录下的文件，获取文件名，并且将文件名后缀修改为txt
# 在word文档中，生成


#--loop
#逐行分析文本，判断函数的开头，如果是方法函数，那么就将该行复制出来  
# 创建一个新的方法单元格， 在第一列中输入‘方法+序号’；
# 将复制的函数名插入第一行第二列。
# 分析该行内容，抓取修饰符后面的内容，插入“返回值“栏
# 分析该行（）之间的内容，以‘，’为分割符，进行获取，将内容分别换行插入到“输入”栏中
#--loop end

# 定义一个单元格类
class Unit :
    methodName=""
    # "对象名："
    inputValue=()
    outputValue=""
    desc=""

    def __init__(self, **kwargs):
        return super().__init__(**kwargs)

 
# 遍历文件夹
folderPath  = "C:/Users/Administrator/Desktop/ctc"
csFileList = ListFolder.gci(folderPath)

# 得到文件列表
# 1 打开word文档
doc = docx.Document()
doc.add_heading(u'写入测试',0)

# 2 遍历文件列表 ，创建对应的表格  ， 写入word文档
for csFile in csFileList:
    className = csFile[csFile.rfind('/')+1:csFile.rfind('.')]
    #try:
    f = open(csFile,'r',encoding='UTF-8')
    line = f.readline()
    #except:
    #    print("exception --------------------------------> ")
    #    continue
    unitsList = []
    i = 0;
    while line :
        if 'public' in line or 'private' in line or 'void' in line or 'protected' in line:
            if '(' in line and '=' not in line:
                line = line.strip()
                print(i,"->"+line)
                # 获取到了字符串方法名 操作word文档，生成表格
                unit = Unit()
                unit.methodName = line[0:line.find(')')+1]
                unit.desc = ""
                #截取（）中的内容
                print("---------------",line.find('('))
                inputArgs = line[line.find('(')+1:line.find(')')]
                #print(inputArgs)
                t = inputArgs.split(',')
                unit.inputValue = t
                inputStr  = ""
                for arg in t:
                    inputStr += arg.lstrip()+":\n"
                #print("====="+inputStr)
                if line.find("/") >= 0:
                    unit.desc = line[line.find("/")+2:len(line)]
                    #print("desc => " + unit.desc)
                # 添加返回值
                parts = line.split(' ')
                if parts[1].find('(') >= 0:
                    unit.outputValue = 'void'
                elif parts[1].find('override')>= 0:
                    unit.outputValue = parts[2]
                else:
                    unit.outputValue = parts[1]
                unitsList.append(unit)

        line = f.readline()
        i += 1
    # 将获取出的注释内容写到word中去
    if unitsList != None:
        print("========================\n",len(unitsList))
        rowCount = len(unitsList) * 4

        # 遍历文件列表，依次增加表格
        # 没个文件，创建一个表格，表头上注明类名和描述
        table = doc.add_table(rows = rowCount+3,cols = 3,style='Table Grid')
        hcells0 = table.rows[0].cells
        hcells0[1].merge(hcells0[2])
        hcells0[0].text = "文件名"
        hcells0[1].text = className
        hcells1 = table.rows[1].cells
        hcells1[1].merge(hcells1[2])
        hcells1[0].text = "描述"
        hcells1[1].text = ""
        hcells2 = table.rows[2].cells
        hcells2[1].merge(hcells2[2])
        hcells2[0].text = "路径"
        hcells2[1].text = csFile.replace(".txt",".cs")
        table.columns[0].width = Inches(0.6)
        table.columns[1].width = Inches(0.6)
        table.columns[2].width = Inches(4.8)

        for i in range(len(unitsList)):
            cells0 = table.rows[(0 + i*4)+3].cells
            
            # TODO 添加方法序号（拼接字符串和整数）2.合并单元格。

            cells0[1].text = "名称："
            cells0[2].text = unitsList[i].methodName
            cells1 = table.rows[(1+i*4)+3].cells
            cells1[1].text = "功能："
            cells1[2].text = unitsList[i].desc
            cells2 = table.rows[(2+i*4)+3].cells
            cells2[1].text = "输入："
            input = ""
            for arg1 in unitsList[i].inputValue:
                if arg1 == None or len(arg1) <= 0:
                    input = ""
                    break;
                input += arg1.lstrip()+" \n"
            input = input[0:len(input)-1]
            print(i,"---"+ input)
            cells2[2].text = input
            cells3 = table.rows[(3+i*4)+3].cells
            cells3[1].text = "输出："
            cells3[2].text = unitsList[i].outputValue
            cells0[0].merge(cells3[0])
            cells0[0].text = "方法{0}".format(i+1)
        doc.add_paragraph(u'\n')
        
    f.close()
# 3 保存文档并关闭
doc.save(u'test2.docx')





