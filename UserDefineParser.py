#coding=utf-8
#!/usr/bin/python3
import docx
import sys
import string
import os
from docx.shared import Inches
import GStack
from GStack import GStack
# 解析class类中的属性
# 1、包含public 和‘；’，则从下标为7的字符截取到；
# 2、包含public 和 ‘}’，截取下标7 到 长度-14的字符串。
# 3、枚举属性，从开头直接到‘，’下标。


# 解析结构体
# 包含struct，截取[struct+8,len)
#  
# 枚举，包含enum ，[8,
# class类中的属性，包含 {get,set}


class FieldUnit:
    fType = ""
    fName = ""
    fDesc = ""
    remark = ""

class ClassUnit:
    className = ""
    classDesc = ""
    fields = []
    def __init__(self, **kwargs):
        return super().__init__(**kwargs)


f = open("C:/Users/Administrator/Desktop/UserDefinitions.txt",'r',encoding='UTF-8')
line = f.readline()

doc = docx.Document()
doc.add_paragraph(u'\n')

className = ""
isInClass = False
isInEnum = False
isScanningProperty = False
gstack = GStack()

classes =  []
classUnit = ClassUnit()
while line:
    line = line.replace("\n","")
    # 获取类名，枚举名 ，结构体名
    if not isInEnum:
        # 枚举
        if ' enum ' in line :
            isInEnum = True
            if '//' in line:
                classUnit.className = line[line.find('enum')+5:line.find("/")]
                classUnit.classDesc = line[line.find("/")+2:len(line)]
            else:
                classUnit.className = line[line.find('enum')+5:line.rfind('\n')]
            print("enum -> "+classUnit.className)
            line = f.readline()
            continue
    else:
        if '{'  in line:
            gstack.push('{')
        if '}' in line:
            gstack.pop()
        if gstack.size() == 0:
            isInEnum = False
            
            rowCount = len(classUnit.fields)
            table = doc.add_table(rows = rowCount + 3,cols = 3,style='Table Grid')
            nameCells = table.rows[0].cells
            nameCells[0].merge(nameCells[2])
            nameCells[0].text  =  "类名："+classUnit.className
            desCells = table.rows[1].cells
            desCells[0].merge(desCells[2])
            desCells[0].text = "描述："+ classUnit.classDesc
            tCells = table.rows[2].cells
            tCells[0].text = "枚举值"
            tCells[1].text = "描述"
            tCells[2].text = "备注"
            table.columns[0].width = Inches(2.4)
            table.columns[1].width = Inches(2.8)
            table.columns[2].width = Inches(0.8)

            for i in range(rowCount):
                fieldCells = table.rows[i+3].cells
                fieldCells[0].text = classUnit.fields[i].fType + " " + classUnit.fields[i].fName
                fieldCells[1].text = classUnit.fields[i].fDesc
                fieldCells[2].text = ""
                print("field->"+ classUnit.className+ "des ->"+classUnit.classDesc)
            doc.add_paragraph(u'\n')
        if gstack.size() == 1:# 目前处于扫描enum属性中
            stripLine = line.strip()
            propLine = stripLine[0:stripLine.find('/')]
            stripPropLine = propLine.strip()
            partsLine = stripPropLine[0:stripPropLine.rfind(',')]
            parts = partsLine.split(',')
            for part in parts:
                field = FieldUnit()
                field.fType = ""
                if '//' in line :
                    field.fDesc = line[line.find('//')+2:len(line)]
                    field.fName = part
                else:
                    field.fName = part
                print("field -> " + field.fName)
                classUnit.fields.append(field)
    if not isInClass:
        #类
        if 'class ' in line:
            if '//' in line:
                classUnit.className = line[line.find("class")+5:line.find("/")]
                classUnit.classDesc = line[line.find("/")+2:len(line)]
            else:
                classUnit.className = line[line.find("class")+5:line.rfind('\n')]
                classUnit.classDesc = ""
            isInClass = True
            print("---"+classUnit.className + "des ->" + classUnit.classDesc)
            line = f.readline()
            continue
        # 结构体 
        
    else:#  在类里
        if '{'  in line:
            gstack.push('{')
        if '}' in line:
            gstack.pop()
        if gstack.size() == 0:
            isInClass = False
            #classes.append(classUnit)
            rowCount = len(classUnit.fields)
            table = doc.add_table(rows = rowCount + 3,cols = 3,style='Table Grid')
            nameCells = table.rows[0].cells
            nameCells[0].merge(nameCells[2])
            nameCells[0].text  =  "类名：" + classUnit.className
            desCells = table.rows[1].cells
            desCells[0].merge(desCells[2])
            desCells[0].text = "描述：" + classUnit.classDesc
            tCells = table.rows[2].cells
            tCells[0].text = "成员属性"
            tCells[1].text = "描述"
            tCells[2].text = "备注"
            table.columns[0].width = Inches(2.4)
            table.columns[1].width = Inches(2.8)
            table.columns[2].width = Inches(0.8)

            for i in range(rowCount):
                fieldCells = table.rows[i+3].cells
                fieldCells[0].text = classUnit.fields[i].fType + " " + classUnit.fields[i].fName
                fieldCells[1].text = classUnit.fields[i].fDesc
                fieldCells[2].text = ""
                print("field->"+ classUnit.className+ "des ->"+classUnit.classDesc)
            doc.add_paragraph(u'\n')

            classUnit.fields = []
            continue
        if gstack.size() == 1:# 目前处于扫描属性中
            if 'public' in line and ';' in line :
                stripLine = line.strip()
                if 'get;' in line and 'set;'in line and '{' in line and '}' in line:
                    propLine = stripLine[0:stripLine.find("{")]
                else :
                    propLine = stripLine[0:stripLine.find(";")]
                #propLine = line.strip()[0:line.find(';')]
                print("propLine-> " + propLine)
                if '= ' in propLine:
                   parts = propLine.strip().split(' ')[0:3]
                else:
                    parts = propLine.strip().split(' ')
                Type = parts[1]
                for part in parts[2:len(parts)]:
                    field = FieldUnit()
                    field.fType = Type
                    stripPart =  part.strip()
                    if '//' in line:
                        field.fDesc = line[line.find('//')+2:len(line)]
                        field.fName = stripPart[0:part.rfind(';')]
                    elif ',' in part :
                       stripPart =  part.strip()
                       field.fName = stripPart[0:part.rfind(',')]
                    elif ';' in line:
                        field.fName = stripPart[0:part.rfind(';')]
                    else :
                        field.fName = part
                    print("field -> " + field.fName)
                    classUnit.fields.append(field)
            if 'public' not in line and 'protected' not in line and 'private' not in line and ';' in line:
                parts = line.strip()[0:line.rfind(';')].split(' ')
                Type = parts[0]
                for  part in parts[2:len(parts)]:
                    field = FieldUnit()
                    field.fType = Type
                    if '//' in line:
                        field.fDesc = line[line.find('//')+2:len(line)]
                        field.fName = part[0:part.rfind(';')]
                    elif ',' in line:
                        field.fName = part.strip()[0:part.rfind(',')]
                    else:
                        field.fName = part
                    print("field -> " + field.fName)
                    classUnit.fields.append(field)
    line = f.readline()
print("class count -> ",len(classes))
f.close()
doc.save(u'user.docx')


